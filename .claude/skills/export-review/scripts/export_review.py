#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""검수용 docx 생성기 (CRATA).

원본 md를 '블록'으로 쪼개, 블록마다 2열 표를 만든다:
  항목 / 위치ID / 원본 / 수정사항(빈칸) / 근거(빈칸)
대표님은 '수정사항'·'근거' 오른쪽 칸만 채운다. 위치ID·원본은 그대로 둔다.
반영은 apply 단계의 parse_review.py가 그 칸을 읽어 처리한다.

사용:
  python export_review.py <source.md> [--out 검수] [--unit auto|heading|bullet|master] [--level N]

블록 단위(--unit):
  master  : '연결키: `...`' 단위. 원본 = ```text``` 본문만(수정 가능한 부분).
  heading : 지정 레벨(--level, 기본=최대 깊이) 헤딩 단위.
  bullet  : '- **라벨**: ...' 최상위 불릿 단위.
  auto    : 연결키 있으면 master, 아니면 heading.
"""
import sys
import re
import argparse
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HEADING_RE = re.compile(r"^(#{1,6})\s+(.*\S)\s*$")
KEY_RE = re.compile(r"^연결키:\s*`([^`]+)`")
FENCE_OPEN_RE = re.compile(r"^```text\s*$")
FENCE_CLOSE_RE = re.compile(r"^```\s*$")
BULLET_RE = re.compile(r"^- \*\*(.+?)\*\*\s*[:：]?\s*(.*)$")
H4_RE = re.compile(r"^####\s+(.*\S)\s*$")


def extract_heading(lines, level=None):
    headings = [(i, len(m.group(1)), m.group(2))
                for i, ln in enumerate(lines)
                for m in [HEADING_RE.match(ln)] if m]
    if not headings:
        return []
    if level is None:
        level = max(h[1] for h in headings)
    blocks = []
    for idx, (line_i, lvl, text) in enumerate(headings):
        if lvl != level:
            continue
        # 조상 경로 (문서 제목 h1은 제외)
        path = []
        prev = lvl
        for j in range(idx - 1, -1, -1):
            hl = headings[j][1]
            if hl < prev:
                if hl == 1:
                    break
                path.insert(0, headings[j][2])
                prev = hl
        path.append(text)
        anchor = " > ".join(path)
        # 본문: 다음 동급/상위 헤딩 전까지
        end = len(lines)
        for j in range(idx + 1, len(headings)):
            if headings[j][1] <= lvl:
                end = headings[j][0]
                break
        body = "\n".join(lines[line_i + 1:end]).strip("\n")
        blocks.append({"title": text, "anchor": anchor, "original": body})
    return blocks


def extract_bullet(lines):
    blocks = []
    i = 0
    n = len(lines)
    while i < n:
        m = BULLET_RE.match(lines[i])
        if m:
            label = m.group(1).strip()
            buf = [lines[i]]
            j = i + 1
            while j < n and (lines[j].startswith("  ") or lines[j].startswith("\t")):
                buf.append(lines[j])
                j += 1
            blocks.append({"title": label, "anchor": label,
                           "original": "\n".join(buf).strip("\n")})
            i = j
        else:
            i += 1
    return blocks


def extract_master(lines):
    blocks = []
    n = len(lines)
    i = 0
    last_h4 = None
    while i < n:
        h4 = H4_RE.match(lines[i])
        if h4:
            last_h4 = h4.group(1)
        m = KEY_RE.match(lines[i])
        if m:
            key = m.group(1)
            # text 블록 찾기
            j = i + 1
            inner = []
            while j < n and not KEY_RE.match(lines[j]):
                if FENCE_OPEN_RE.match(lines[j]):
                    j += 1
                    while j < n and not FENCE_CLOSE_RE.match(lines[j]):
                        inner.append(lines[j])
                        j += 1
                    break
                j += 1
            blocks.append({"title": last_h4 or key, "anchor": key,
                           "original": "\n".join(inner).strip("\n")})
            last_h4 = None
        i += 1
    return blocks


def shade(cell, hexcolor):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcpr.append(shd)


AXIS = {"O": "목적성취방식", "P": "능력표현방식", "A": "목표실행방식", "S": "정보처리방식"}


def parse_fields(md):
    """마크다운 본문을 (라벨|None, 값) 목록으로."""
    fields = []
    for raw in md.splitlines():
        s = raw.strip()
        if not s:
            continue
        if re.match(r"^([-_*])\1{2,}$", s):  # 수평선(---) 제거
            continue
        if s.startswith(">"):  # 인용 기호 제거, 본문만
            s = s.lstrip(">").strip()
            if not s:
                continue
        m = re.match(r"^-\s*\*\*(.+?)\*\*\s*[:：]\s*(.*)$", s)
        if m:
            fields.append((m.group(1).strip(), m.group(2).strip()))
            continue
        m2 = re.match(r"^-\s*\*\*(.+?)\*\*\s*$", s)
        if m2:
            fields.append((m2.group(1).strip(), ""))
            continue
        m3 = re.match(r"^-\s+(.*)$", s)
        if m3:
            fields.append((None, m3.group(1).strip()))
            continue
        fields.append((None, s))
    return fields


def clean(text):
    return re.sub(r"\*+", "", text)


def pretty_type_title(title):
    """'OA 책임실행형 (1010 · O+A)' → 'OA 책임실행형 (목적성취방식 + 목표실행방식)'."""
    m = re.match(r"^(.*?)\s*\(", title)
    ax = re.search(r"·\s*([^)]+)\)", title)
    if not m or not ax:
        return title
    name = m.group(1).strip()
    tok = ax.group(1).strip()
    letters = re.findall(r"[OPAS]", tok)
    if letters:
        names = " + ".join(AXIS[c] for c in letters)
    else:
        names = "해당 방식 없음" if "없음" in tok else tok
    return f"{name} ({names})"


def display_title(blk):
    anchor = blk["anchor"]
    if ">" in anchor:
        return " > ".join(p.strip() for p in anchor.split(">"))
    return pretty_type_title(blk["title"])


def fill_title(cell, blk):
    cell.text = ""
    run = cell.paragraphs[0].add_run(display_title(blk))
    run.font.bold = True
    run.font.size = Pt(10)


def fill_original(cell, blk):
    rendered = []
    title = blk["title"]
    for label, val in parse_fields(blk["original"]):
        if label == "소속 방식":
            continue  # 계층 헤더에 이미 나타남
        if label and (label == title or title.startswith(label)):
            label = "정의"  # 불릿 값은 곧 그 유형의 정의
        rendered.append((label, clean(val)))
    cell.text = ""
    first = True
    for label, val in rendered:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        if label:
            r1 = p.add_run(f"{label}: ")
            r1.font.bold = True
            r1.font.size = Pt(10)
        r2 = p.add_run(val)
        r2.font.size = Pt(10)


def add_block_table(doc, blk):
    labels = ["항목", "위치ID", "원본", "수정사항", "근거"]
    table = doc.add_table(rows=len(labels), cols=2)
    table.style = "Table Grid"
    for r, label in enumerate(labels):
        lc = table.rows[r].cells[0]
        vc = table.rows[r].cells[1]
        lc.text = label
        lc.paragraphs[0].runs[0].font.bold = True
        lc.paragraphs[0].runs[0].font.size = Pt(9)
        if label == "항목":
            fill_title(vc, blk)
            shade(lc, "EFEFEF")
        elif label == "위치ID":
            vc.text = blk["anchor"]
            for run in vc.paragraphs[0].runs:
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            shade(lc, "EFEFEF")
        elif label == "원본":
            fill_original(vc, blk)
            shade(lc, "EFEFEF")
        else:  # 수정사항, 근거: 입력 칸(옅은 노랑)
            vc.text = ""
            shade(lc, "FFF2CC")
    for row in table.rows:
        row.cells[0].width = Pt(70)
    doc.add_paragraph("")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("source")
    ap.add_argument("--out", default="검수")
    ap.add_argument("--unit", default="auto",
                    choices=["auto", "heading", "bullet", "master"])
    ap.add_argument("--level", type=int, default=None)
    args = ap.parse_args()

    src = Path(args.source)
    lines = src.read_text(encoding="utf-8").splitlines()

    unit = args.unit
    if unit == "auto":
        unit = "master" if any(KEY_RE.match(l) for l in lines) else "heading"

    if unit == "master":
        blocks = extract_master(lines)
    elif unit == "bullet":
        blocks = extract_bullet(lines)
    else:
        blocks = extract_heading(lines, args.level)

    blocks = [b for b in blocks if b["original"].strip()]
    if not blocks:
        print("블록을 찾지 못함. --unit/--level 확인.", file=sys.stderr)
        return 2

    doc = Document()
    doc.add_heading(f"CRATA 검수 — {src.stem}", level=0)
    p = doc.add_paragraph()
    p.add_run("작성 후 ").font.size = Pt(10)
    r = p.add_run("‘수정사항’·‘근거’ 노란 칸만 채워주세요.")
    r.font.bold = True
    r.font.size = Pt(10)
    p.add_run(" ‘위치ID’·‘원본’은 그대로 두세요(반영 시 위치 매칭에 씁니다).").font.size = Pt(10)
    doc.add_paragraph(f"원본 파일: {src.as_posix()}    |    단위: {unit}    |    생성일: {date.today().isoformat()}")
    doc.add_paragraph("")

    for blk in blocks:
        add_block_table(doc, blk)

    out_dir = Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"{date.today().isoformat()}_{src.stem}_검수.docx"
    doc.save(str(out_path))
    print(f"생성: {out_path}  (블록 {len(blocks)}개, 단위 {unit})")
    return 0


if __name__ == "__main__":
    sys.exit(main())
