#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""채워진 검수 docx 파서 (CRATA).

export_review.py가 만든 docx를 읽어, '수정사항'이 채워진 블록만 JSON으로 출력한다.
출력(JSON 배열, 각 원소):
  {title, anchor, original, revision, reason}
apply 단계(Claude)는 이 JSON을 보고 anchor로 원본 파일 위치를 찾아
original→revision으로 제안→확인→반영하고, reason을 decisions/log에 남긴다.

사용:
  python parse_review.py <filled.docx> [--all]
  --all : 수정사항이 빈 블록까지 전부 출력(점검용).
"""
import sys
import json
import argparse
from pathlib import Path

from docx import Document

LABELS = ("항목", "위치ID", "원본", "수정사항", "근거")


def cell_text(cell):
    return "\n".join(p.text for p in cell.paragraphs).strip()


def parse_table(table):
    data = {}
    for row in table.rows:
        if len(row.cells) < 2:
            continue
        label = cell_text(row.cells[0])
        if label in LABELS:
            data[label] = cell_text(row.cells[1])
    if "위치ID" not in data and "원본" not in data:
        return None
    return {
        "title": data.get("항목", ""),
        "anchor": data.get("위치ID", ""),
        "original": data.get("원본", ""),
        "revision": data.get("수정사항", ""),
        "reason": data.get("근거", ""),
    }


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("docx")
    ap.add_argument("--all", action="store_true")
    args = ap.parse_args()

    doc = Document(args.docx)
    out = []
    for table in doc.tables:
        rec = parse_table(table)
        if rec is None:
            continue
        if args.all or rec["revision"].strip():
            out.append(rec)

    sys.stdout.reconfigure(encoding="utf-8")
    print(json.dumps(out, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    sys.exit(main())
